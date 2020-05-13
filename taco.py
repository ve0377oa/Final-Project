'''Mohamed Mohamed May 13, 2020  ITEC 1150   Final Project'''

# Use requests to download three random tacos from the random taco API.
# Save the data for each of three tacos in your program.
# Notice that each recipe is divided into five sections for
# base_layer,
# seasoning,
# mixin,
# condiment,
# and shell.
# Use Python to create a Word document.
# On the first page, insert the header text "Random Taco Cookbook"
# On the first page, add the resized taco image that you created with
#
# part 1. (hint: adding images to Word documents is covered in the textbook)
# On the first page, write the name of the image author
# On the first page, write the text of the random taco API URL
# On the first page, write your own name.
#
# On the second page, start writing the first taco recipe.
# Write all five components of the recipe.
# Use a larger font or heading for the heading for each of the sections.
# Please see example document for suggested style.
# After the first recipe, add a page break.
# To add another page, hint: google "python-docx add page break"
# Repeat to write all of the next recipe and a page break.
# Repeat to write all the third recipe.
# Save your word document.


from docx import Document
from docx.shared import Inches
import requests

###########################################
# First Page
###########################################

# constants
IMAGE_AUTHOR = "Image Author: Tai's Captures"
TACO_API_URL = "https://taco-1150.herokuapp.com/random/?full_taco=true"
STUDENT_NAME = "Student Name: DQ"

document = Document()  # object created

document.add_heading('Random Taco Cookbook', 0)  # add heading
document.add_picture('taco_original_thumbnail.jpg', width=Inches(6.67))  # add picture
document.add_paragraph(IMAGE_AUTHOR)  # add text
document.add_paragraph(TACO_API_URL)  # add text
document.add_paragraph(STUDENT_NAME)  # add text


###########################################
# Recipes
###########################################

# fetchJSON
# get JSON response
# params:
#   url:takes taco api url
# return:
#   get_json:json containing taco recipe
def fetchJSON(url):
    response = requests.get(url)  # get response from taco url api
    get_json = response.json()  # get json
    return get_json


# addRecipe
# navigates json as dictinary and fetches recipe
# params:
#   json_file: takes json file for parsing
#   document: takes document object
def addRecipe(json_file, document):
    for component in json_file:
        p = document.add_heading(component, level=1)  # add heading as component of recipe
        for item in json_file[component]:
            if item == 'name':
                name = json_file[component][item]
                document.add_heading(name, level=1)  # add heading as recipe component's name
            if item == 'recipe':
                recipe = json_file[component][item]
                document.add_paragraph(recipe)  # add text as recipe


for i in range(3):
    document.add_page_break()
    document.add_heading(f'Random Taco Recipe {i + 1}', 0)  # add heading for recipes
    addRecipe(fetchJSON(TACO_API_URL), document)

document.save('Final Project.docx')  # save document